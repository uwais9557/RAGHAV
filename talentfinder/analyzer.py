"""
Talent Finder - Resume Analysis Engine
=======================================
AI-powered resume text extraction, skill detection, and job matching
using NLP, TF-IDF, and cosine similarity.
"""

import re
import json
import logging
from typing import List, Dict, Tuple

logger = logging.getLogger(__name__)

# ─── Skill Database ────────────────────────────────────────────────────────────
SKILL_KEYWORDS = {
    # Programming Languages
    'python': 'programming', 'java': 'programming', 'javascript': 'programming',
    'c++': 'programming', 'c#': 'programming', 'php': 'programming',
    'ruby': 'programming', 'swift': 'programming', 'kotlin': 'programming',
    'go': 'programming', 'rust': 'programming', 'scala': 'programming',
    'r': 'programming', 'matlab': 'programming', 'typescript': 'programming',

    # Web Technologies
    'html': 'web', 'css': 'web', 'html5': 'web', 'css3': 'web',
    'react': 'web', 'angular': 'web', 'vue': 'web', 'nodejs': 'web',
    'bootstrap': 'web', 'jquery': 'web', 'sass': 'web', 'webpack': 'web',
    'rest api': 'web', 'graphql': 'web', 'next.js': 'web', 'flask': 'web',
    'django': 'framework', 'spring boot': 'framework', 'express': 'framework',

    # Databases
    'sql': 'database', 'mysql': 'database', 'postgresql': 'database',
    'mongodb': 'database', 'redis': 'database', 'sqlite': 'database',
    'oracle': 'database', 'firebase': 'database', 'elasticsearch': 'database',
    'cassandra': 'database', 'dynamodb': 'database',

    # Data Science & AI/ML
    'machine learning': 'data_science', 'deep learning': 'data_science',
    'artificial intelligence': 'data_science', 'data science': 'data_science',
    'tensorflow': 'data_science', 'pytorch': 'data_science', 'keras': 'data_science',
    'scikit-learn': 'data_science', 'pandas': 'data_science', 'numpy': 'data_science',
    'nlp': 'data_science', 'computer vision': 'data_science',
    'neural networks': 'data_science', 'data analysis': 'data_science',
    'data visualization': 'data_science', 'statistics': 'data_science',
    'power bi': 'tools', 'tableau': 'tools', 'excel': 'tools',

    # Cloud & DevOps
    'aws': 'cloud', 'azure': 'cloud', 'google cloud': 'cloud',
    'docker': 'cloud', 'kubernetes': 'cloud', 'jenkins': 'cloud',
    'git': 'tools', 'github': 'tools', 'gitlab': 'tools',
    'linux': 'tools', 'devops': 'cloud', 'ci/cd': 'cloud',
    'terraform': 'cloud', 'ansible': 'cloud',

    # Mobile
    'android': 'programming', 'ios': 'programming', 'flutter': 'framework',
    'react native': 'framework', 'xamarin': 'framework',

    # Tools
    'jira': 'tools', 'agile': 'tools', 'scrum': 'tools',
    'photoshop': 'tools', 'figma': 'tools', 'postman': 'tools',
    'selenium': 'tools', 'junit': 'tools',
}


# ─── Text Extraction ───────────────────────────────────────────────────────────
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using PyPDF2"""
    try:
        import PyPDF2
        text = ""
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return clean_text(text)
    except ImportError:
        logger.error("PyPDF2 not installed")
        return ""
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx"""
    try:
        from docx import Document
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return clean_text(text)
    except ImportError:
        logger.error("python-docx not installed")
        return ""
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return ""


def extract_text_from_image(file_path: str) -> str:
    """Extract text from image using pytesseract OCR"""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img, lang='eng')
        return clean_text(text)
    except ImportError:
        logger.error("pytesseract or Pillow not installed")
        return "OCR extraction requires pytesseract installation"
    except Exception as e:
        logger.error(f"Image OCR error: {e}")
        return ""


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return clean_text(f.read())
    except Exception as e:
        logger.error(f"TXT extraction error: {e}")
        return ""


def extract_text(file_path: str, file_type: str) -> str:
    """Master function to extract text based on file type"""
    extractors = {
        'pdf': extract_text_from_pdf,
        'docx': extract_text_from_docx,
        'doc': extract_text_from_docx,
        'txt': extract_text_from_txt,
        'png': extract_text_from_image,
        'jpg': extract_text_from_image,
        'jpeg': extract_text_from_image,
    }
    extractor = extractors.get(file_type.lower(), extract_text_from_txt)
    return extractor(file_path)


def clean_text(text: str) -> str:
    """Clean and normalize extracted text"""
    if not text:
        return ""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters but keep important ones
    text = re.sub(r'[^\w\s\.\,\@\+\-\/\(\)\#]', ' ', text)
    return text.strip()


# ─── Information Extraction ────────────────────────────────────────────────────
def extract_email(text: str) -> str:
    """Extract email address from text"""
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    matches = re.findall(pattern, text)
    return matches[0] if matches else ""


def extract_phone(text: str) -> str:
    """Extract phone number from text"""
    pattern = r'(\+91[\-\s]?)?[6-9]\d{9}|(\+\d{1,3}[\-\s]?)?\d{3}[\-\s]?\d{3}[\-\s]?\d{4}'
    matches = re.findall(pattern, text)
    if matches:
        return ''.join(matches[0]).strip()
    return ""


def extract_candidate_name(text: str) -> str:
    """Attempt to extract candidate name from first lines of resume"""
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines[:5]:
        words = line.split()
        if 2 <= len(words) <= 4:
            if all(w[0].isupper() for w in words if w):
                if not any(k in line.lower() for k in ['resume', 'cv', 'curriculum', 'email', 'phone']):
                    return line
    return ""


def detect_experience_years(text: str) -> float:
    """Detect years of experience from text"""
    text_lower = text.lower()
    patterns = [
        r'(\d+\.?\d*)\s*(?:\+\s*)?years?\s*(?:of\s*)?(?:work\s*)?experience',
        r'experience\s*(?:of\s*)?(\d+\.?\d*)\s*years?',
        r'(\d+\.?\d*)\s*years?\s*in\s*(?:the\s*)?(?:industry|field|it|software)',
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text_lower)
        if matches:
            try:
                return float(matches[0])
            except (ValueError, IndexError):
                pass
    return 0.0


def detect_education_level(text: str) -> str:
    """Detect highest education level from text"""
    text_lower = text.lower()
    levels = [
        ('phd', 'PhD / Doctorate'),
        ('doctorate', 'PhD / Doctorate'),
        ('m.tech', 'M.Tech'),
        ('m.e.', 'M.E.'),
        ('mca', 'MCA'),
        ('mba', 'MBA'),
        ('master', 'Master\'s Degree'),
        ('b.tech', 'B.Tech'),
        ('b.e.', 'B.E.'),
        ('bca', 'BCA'),
        ('bachelor', 'Bachelor\'s Degree'),
        ('diploma', 'Diploma'),
        ('12th', '12th Grade'),
        ('10th', '10th Grade'),
    ]
    for keyword, display in levels:
        if keyword in text_lower:
            return display
    return "Not Detected"


# ─── Skill Detection ───────────────────────────────────────────────────────────
def detect_skills(text: str) -> List[str]:
    """Detect skills from resume text using keyword matching"""
    text_lower = text.lower()
    found_skills = []

    for skill, category in SKILL_KEYWORDS.items():
        # Create word boundary pattern
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            found_skills.append(skill.title())

    return list(set(found_skills))


# ─── ATS Score Calculation ─────────────────────────────────────────────────────
def calculate_ats_score(text: str, skills: List[str]) -> int:
    """
    Calculate ATS (Applicant Tracking System) score based on:
    - Number of detected skills (40%)
    - Contact information presence (20%)
    - Education information (15%)
    - Experience information (15%)
    - Resume length and structure (10%)
    """
    score = 0

    # Skills score (max 40 points)
    skill_score = min(len(skills) * 3, 40)
    score += skill_score

    # Contact info (max 20 points)
    if extract_email(text):
        score += 10
    if extract_phone(text):
        score += 10

    # Education (max 15 points)
    education = detect_education_level(text)
    if education not in ['Not Detected', '10th Grade']:
        score += 15
    elif education == '10th Grade':
        score += 5

    # Experience (max 15 points)
    exp = detect_experience_years(text)
    if exp > 5:
        score += 15
    elif exp > 2:
        score += 10
    elif exp > 0:
        score += 7
    else:
        score += 3  # Fresher gets some points

    # Resume structure/length (max 10 points)
    word_count = len(text.split())
    if word_count > 300:
        score += 10
    elif word_count > 150:
        score += 7
    elif word_count > 50:
        score += 4

    return min(score, 100)


# ─── Job Matching ──────────────────────────────────────────────────────────────
def calculate_job_match(resume_skills: List[str], job_required_skills: List[str]) -> Tuple[float, List[str], List[str]]:
    """
    Calculate match percentage between resume skills and job requirements
    using set intersection and cosine-like similarity.

    Returns: (match_percentage, matched_skills, missing_skills)
    """
    if not job_required_skills:
        return 0.0, [], []

    resume_set = set(s.lower() for s in resume_skills)
    job_set = set(s.lower() for s in job_required_skills)

    matched = resume_set.intersection(job_set)
    missing = job_set - resume_set

    if not job_set:
        return 0.0, [], []

    # Weighted similarity: matched / total required
    match_pct = (len(matched) / len(job_set)) * 100

    # Convert back to proper case
    matched_display = [s.title() for s in matched]
    missing_display = [s.title() for s in missing]

    return round(match_pct, 1), matched_display, missing_display


def get_tfidf_similarity(resume_text: str, job_description: str) -> float:
    """Calculate TF-IDF cosine similarity between resume and job description"""
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity

        if not resume_text or not job_description:
            return 0.0

        vectorizer = TfidfVectorizer(stop_words='english', ngram_range=(1, 2))
        tfidf_matrix = vectorizer.fit_transform([resume_text, job_description])
        similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
        return round(float(similarity[0][0]) * 100, 1)
    except ImportError:
        return 0.0
    except Exception as e:
        logger.error(f"TF-IDF similarity error: {e}")
        return 0.0


def recommend_jobs(resume, limit: int = 5) -> List[Dict]:
    """
    Recommend top jobs based on resume skills using hybrid matching.
    Combines skill-based matching with TF-IDF similarity.
    """
    from .models import Job, JobMatch

    resume_skill_names = [s.name.lower() for s in resume.detected_skills.all()]
    jobs = Job.objects.filter(is_active=True).prefetch_related('required_skills', 'company')

    job_scores = []
    for job in jobs:
        job_skill_names = [s.name.lower() for s in job.required_skills.all()]
        if not job_skill_names:
            continue

        # Skill-based match
        skill_match, matched, missing = calculate_job_match(resume_skill_names, job_skill_names)

        # TF-IDF bonus (max 20% bonus)
        tfidf_score = get_tfidf_similarity(resume.extracted_text, job.description)
        tfidf_bonus = tfidf_score * 0.2

        # Combined score
        final_score = min(skill_match + tfidf_bonus, 100)

        job_scores.append({
            'job': job,
            'match_percentage': final_score,
            'matched_skills': matched,
            'missing_skills': missing,
        })

    # Sort by match percentage
    job_scores.sort(key=lambda x: x['match_percentage'], reverse=True)

    # Save top matches to database
    JobMatch.objects.filter(resume=resume).delete()
    for item in job_scores[:limit]:
        jm = JobMatch.objects.create(
            resume=resume,
            job=item['job'],
            match_percentage=item['match_percentage'],
            missing_skills=json.dumps(item['missing_skills']),
        )
        matched_skill_objs = resume.detected_skills.filter(
            name__in=[s.title() for s in item['matched_skills']]
        )
        jm.matched_skills.set(matched_skill_objs)

    return job_scores[:limit]


def suggest_courses(missing_skills: List[str]) -> List[Dict]:
    """Suggest online courses for missing skills"""
    course_map = {
        'python': {'name': 'Python for Everybody', 'platform': 'Coursera', 'url': 'https://coursera.org', 'icon': 'fa-python'},
        'machine learning': {'name': 'Machine Learning Specialization', 'platform': 'Coursera', 'url': 'https://coursera.org', 'icon': 'fa-brain'},
        'deep learning': {'name': 'Deep Learning Specialization', 'platform': 'Coursera', 'url': 'https://coursera.org', 'icon': 'fa-microchip'},
        'data science': {'name': 'Data Science Professional Certificate', 'platform': 'edX', 'url': 'https://edx.org', 'icon': 'fa-chart-bar'},
        'react': {'name': 'React - The Complete Guide', 'platform': 'Udemy', 'url': 'https://udemy.com', 'icon': 'fa-atom'},
        'django': {'name': 'Django for Beginners', 'platform': 'Udemy', 'url': 'https://udemy.com', 'icon': 'fa-server'},
        'sql': {'name': 'SQL Bootcamp', 'platform': 'Udemy', 'url': 'https://udemy.com', 'icon': 'fa-database'},
        'aws': {'name': 'AWS Certified Solutions Architect', 'platform': 'AWS', 'url': 'https://aws.amazon.com/training/', 'icon': 'fa-cloud'},
        'docker': {'name': 'Docker Mastery', 'platform': 'Udemy', 'url': 'https://udemy.com', 'icon': 'fa-box'},
        'kubernetes': {'name': 'Kubernetes Complete Guide', 'platform': 'Udemy', 'url': 'https://udemy.com', 'icon': 'fa-dharmachakra'},
        'java': {'name': 'Java Programming Masterclass', 'platform': 'Udemy', 'url': 'https://udemy.com', 'icon': 'fa-coffee'},
        'javascript': {'name': 'JavaScript: The Complete Guide', 'platform': 'Udemy', 'url': 'https://udemy.com', 'icon': 'fa-js'},
        'power bi': {'name': 'Power BI Data Analyst', 'platform': 'Microsoft', 'url': 'https://learn.microsoft.com', 'icon': 'fa-chart-pie'},
        'tensorflow': {'name': 'TensorFlow Developer Certificate', 'platform': 'Google', 'url': 'https://www.tensorflow.org/certificate', 'icon': 'fa-brain'},
        'android': {'name': 'Android Development with Kotlin', 'platform': 'Google', 'url': 'https://developer.android.com/courses', 'icon': 'fa-mobile'},
        'git': {'name': 'Git & GitHub Complete Guide', 'platform': 'Udemy', 'url': 'https://udemy.com', 'icon': 'fa-code-branch'},
    }

    suggestions = []
    for skill in missing_skills[:6]:
        skill_lower = skill.lower()
        if skill_lower in course_map:
            course = course_map[skill_lower].copy()
            course['skill'] = skill
            suggestions.append(course)
        else:
            suggestions.append({
                'skill': skill,
                'name': f'Learn {skill}',
                'platform': 'Udemy',
                'url': f'https://udemy.com/courses/search/?q={skill.replace(" ", "+")}',
                'icon': 'fa-graduation-cap'
            })
    return suggestions
